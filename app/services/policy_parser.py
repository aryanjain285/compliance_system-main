"""
Policy Document Parser Service
Handles PDF parsing, text extraction, and document processing
"""
import os
import re
import json
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from datetime import datetime
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
from sqlalchemy.orm import Session

from app.models.database import PolicyDocument, PolicyChunk, DocumentStatus
from app.utils.logger import get_logger, log_execution_time
from app.utils.exceptions import (
    PolicyParsingException, FileUploadException, ValidationException,
    UnsupportedFileType, FileTooLarge
)
from app.config.settings import get_settings

settings = get_settings()
logger = get_logger(__name__)


class PolicyParser:
    """Advanced policy document parser with support for multiple file formats"""
    
    def __init__(self, db_session: Session):
        self.db = db_session
        self.supported_types = {
            '.pdf': self._extract_pdf_text,
            '.txt': self._extract_text_file,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.md': self._extract_markdown_file,
            '.xlsx': self._extract_excel_file,
            '.csv': self._extract_csv_file
        }
        
        # Create upload directory
        settings.create_upload_dir()
    
    @log_execution_time("upload_policy_document")
    def upload_policy_document(self, filename: str, content: bytes, 
                             uploaded_by: str = "system", 
                             metadata: Dict[str, Any] = None) -> str:
        """Upload and process a policy document with comprehensive validation"""
        try:
            # Validate file
            self._validate_file(filename, content)
            
            # Calculate content hash for deduplication
            content_hash = hashlib.sha256(content).hexdigest()
            
            # Check for existing document with same hash
            existing_doc = self.db.query(PolicyDocument).filter(
                PolicyDocument.content_hash == content_hash
            ).first()
            
            if existing_doc:
                logger.info(f"Document with same content already exists: {existing_doc.policy_id}")
                return existing_doc.policy_id
            
            # Extract text content
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_types:
                raise UnsupportedFileType(file_ext, list(self.supported_types.keys()))
            
            text_content = self.supported_types[file_ext](content)
            
            if not text_content or len(text_content.strip()) < 50:
                raise PolicyParsingException("Insufficient content extracted from document")
            
            # Detect document metadata
            detected_metadata = self._detect_document_metadata(text_content, filename)
            if metadata:
                detected_metadata.update(metadata)
            
            # Create policy document record
            policy = PolicyDocument(
                filename=filename,
                content=text_content,
                content_hash=content_hash,
                uploaded_by=uploaded_by,
                document_type=detected_metadata.get('document_type'),
                jurisdiction=detected_metadata.get('jurisdiction'),
                effective_date=detected_metadata.get('effective_date'),
                expiry_date=detected_metadata.get('expiry_date'),
                metadata={
                    **detected_metadata,
                    "file_size": len(content),
                    "content_length": len(text_content),
                    "file_extension": file_ext,
                    "mime_type": mimetypes.guess_type(filename)[0]
                }
            )
            
            self.db.add(policy)
            self.db.commit()
            self.db.refresh(policy)
            
            # Chunk the document
            chunks = self._chunk_document(text_content, policy.policy_id)
            
            # Store chunks
            for chunk_data in chunks:
                policy_chunk = PolicyChunk(**chunk_data)
                self.db.add(policy_chunk)
            
            self.db.commit()
            
            logger.info(
                f"Policy document uploaded successfully: {policy.policy_id}, "
                f"extracted {len(text_content)} characters, created {len(chunks)} chunks"
            )
            
            return policy.policy_id
            
        except Exception as e:
            self.db.rollback()
            if isinstance(e, (PolicyParsingException, FileUploadException, ValidationException)):
                raise
            
            logger.error(f"Error uploading policy document {filename}: {e}")
            raise PolicyParsingException(f"Failed to upload policy document: {str(e)}")
    
    def _validate_file(self, filename: str, content: bytes):
        """Validate file size and type"""
        # Check file size
        if len(content) > settings.max_file_size_bytes:
            raise FileTooLarge(len(content), settings.max_file_size_bytes)
        
        # Check file extension
        file_ext = Path(filename).suffix.lower().lstrip('.')
        if file_ext not in settings.allowed_file_types:
            raise UnsupportedFileType(file_ext, settings.allowed_file_types)
        
        # Basic content validation
        if len(content) == 0:
            raise FileUploadException("Empty file uploaded")
        
        # Check for potential malicious content
        if b'<script>' in content.lower() or b'javascript:' in content.lower():
            raise FileUploadException("Potentially malicious content detected")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text content from PDF bytes"""
        try:
            import io
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise PolicyParsingException("PDF contains no pages")
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    text_content.append(f"[Page {page_num + 1}]\n[Text extraction failed]")
            
            extracted_text = "\n\n".join(text_content)
            
            if len(extracted_text.strip()) < 50:
                raise PolicyParsingException("Insufficient text extracted from PDF")
            
            return extracted_text
            
        except PyPDF2.errors.PdfReadError as e:
            raise PolicyParsingException(f"Invalid PDF file: {str(e)}")
        except Exception as e:
            raise PolicyParsingException(f"PDF text extraction error: {str(e)}")
    
    def _extract_text_file(self, content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    # Validate that it's readable text
                    if len(text.strip()) > 0:
                        return self._clean_text_content(text)
                except UnicodeDecodeError:
                    continue
            
            raise PolicyParsingException("Unable to decode text file with supported encodings")
            
        except Exception as e:
            raise PolicyParsingException(f"Text file extraction error: {str(e)}")
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            import io
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_content)
            
            if len(extracted_text.strip()) < 50:
                raise PolicyParsingException("Insufficient text extracted from DOCX")
            
            return self._clean_text_content(extracted_text)
            
        except Exception as e:
            raise PolicyParsingException(f"DOCX extraction error: {str(e)}")
    
    def _extract_doc_text(self, content: bytes) -> str:
        """Extract text from DOC file (legacy format)"""
        try:
            # For DOC files, we need a different approach since python-docx doesn't support .doc
            # This is a placeholder - in production you'd use libraries like python-docx2txt
            # or convert to DOCX first
            raise PolicyParsingException(
                "DOC format not fully supported. Please convert to DOCX format."
            )
            
        except Exception as e:
            raise PolicyParsingException(f"DOC extraction error: {str(e)}")
    
    def _extract_markdown_file(self, content: bytes) -> str:
        """Extract text from Markdown file"""
        try:
            text = content.decode('utf-8')
            
            # Basic markdown cleanup - remove common markdown syntax
            text = re.sub(r'#+\s*', '', text)  # Remove headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove inline code
            text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Remove code blocks
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Remove links
            
            return self._clean_text_content(text)
            
        except Exception as e:
            raise PolicyParsingException(f"Markdown extraction error: {str(e)}")
    
    def _extract_excel_file(self, content: bytes) -> str:
        """Extract text from Excel file"""
        try:
            import io
            excel_file = io.BytesIO(content)
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None, dtype=str)
            
            text_content = []
            
            for sheet_name, df in excel_data.items():
                text_content.append(f"[Sheet: {sheet_name}]")
                
                # Convert dataframe to text
                for idx, row in df.iterrows():
                    row_text = []
                    for col, value in row.items():
                        if pd.notna(value) and str(value).strip():
                            row_text.append(f"{col}: {value}")
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_content)
            
            if len(extracted_text.strip()) < 50:
                raise PolicyParsingException("Insufficient text extracted from Excel")
            
            return self._clean_text_content(extracted_text)
            
        except Exception as e:
            raise PolicyParsingException(f"Excel extraction error: {str(e)}")
    
    def _extract_csv_file(self, content: bytes) -> str:
        """Extract text from CSV file"""
        try:
            import io
            csv_file = io.StringIO(content.decode('utf-8'))
            df = pd.read_csv(csv_file, dtype=str)
            
            text_content = []
            
            # Add headers
            text_content.append("CSV Headers: " + " | ".join(df.columns.tolist()))
            
            # Convert rows to text
            for idx, row in df.iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value) and str(value).strip():
                        row_text.append(f"{col}: {value}")
                if row_text:
                    text_content.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_content)
            
            if len(extracted_text.strip()) < 50:
                raise PolicyParsingException("Insufficient text extracted from CSV")
            
            return self._clean_text_content(extracted_text)
            
        except Exception as e:
            raise PolicyParsingException(f"CSV extraction error: {str(e)}")
    
    def _clean_text_content(self, text: str) -> str:
        """Clean and normalize extracted text content"""
        try:
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove control characters but keep newlines and tabs
            text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
            
            # Normalize line endings
            text = re.sub(r'\r\n|\r', '\n', text)
            
            # Remove excessive blank lines
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
            
            return text.strip()
            
        except Exception as e:
            logger.warning(f"Error cleaning text content: {e}")
            return text
    
    def _detect_document_metadata(self, content: str, filename: str) -> Dict[str, Any]:
        """Detect document metadata from content and filename"""
        metadata = {}
        
        try:
            # Detect document type from filename and content
            filename_lower = filename.lower()
            content_lower = content.lower()
            
            # Document type detection
            if any(term in filename_lower for term in ['policy', 'procedure', 'guideline']):
                metadata['document_type'] = 'Policy Document'
            elif any(term in filename_lower for term in ['compliance', 'regulatory']):
                metadata['document_type'] = 'Compliance Manual'
            elif any(term in filename_lower for term in ['risk', 'control']):
                metadata['document_type'] = 'Risk Management'
            elif any(term in filename_lower for term in ['investment', 'trading']):
                metadata['document_type'] = 'Investment Policy'
            else:
                # Detect from content
                if any(term in content_lower for term in ['policy', 'procedure', 'shall', 'must']):
                    metadata['document_type'] = 'Policy Document'
                else:
                    metadata['document_type'] = 'General Document'
            
            # Jurisdiction detection
            jurisdiction_patterns = {
                'US': ['united states', 'sec ', 'finra', 'cftc', 'federal register'],
                'EU': ['european union', 'mifid', 'esma', 'regulation (eu)'],
                'UK': ['united kingdom', 'fca', 'bank of england', 'pra'],
                'Canada': ['canada', 'csa', 'iiroc', 'osfi'],
                'Australia': ['australia', 'asic', 'apra', 'asx'],
                'Singapore': ['singapore', 'mas', 'monetary authority']
            }
            
            for jurisdiction, patterns in jurisdiction_patterns.items():
                if any(pattern in content_lower for pattern in patterns):
                    metadata['jurisdiction'] = jurisdiction
                    break
            
            # Date extraction
            date_patterns = [
                r'effective\s+(?:date|from)?\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{4})',
                r'effective\s+(?:date|from)?\s*:?\s*(\d{4}[/-]\d{1,2}[/-]\d{1,2})',
                r'dated?\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{4})',
                r'(?:as\s+of|date)\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{4})'
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, content_lower)
                if match:
                    try:
                        date_str = match.group(1)
                        # Try to parse date
                        for fmt in ['%m/%d/%Y', '%d/%m/%Y', '%Y/%m/%d', '%m-%d-%Y', '%d-%m-%Y', '%Y-%m-%d']:
                            try:
                                effective_date = datetime.strptime(date_str, fmt)
                                metadata['effective_date'] = effective_date
                                break
                            except ValueError:
                                continue
                        break
                    except:
                        continue
            
            # Version detection
            version_patterns = [
                r'version\s*:?\s*(\d+(?:\.\d+)*)',
                r'v(\d+(?:\.\d+)*)',
                r'revision\s*:?\s*(\d+)'
            ]
            
            for pattern in version_patterns:
                match = re.search(pattern, content_lower)
                if match:
                    metadata['version'] = match.group(1)
                    break
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error detecting document metadata: {e}")
            return metadata
    
    @log_execution_time("chunk_document")
    def _chunk_document(self, content: str, policy_id: str, 
                       chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """Intelligently chunk document content with overlap and section awareness"""
        try:
            chunks = []
            lines = content.split('\n')
            current_chunk = []
            current_size = 0
            chunk_index = 0
            current_page = 1
            current_section = "Introduction"
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect page breaks
                page_match = re.match(r'\[Page (\d+)\]', line)
                if page_match:
                    current_page = int(page_match.group(1))
                    continue
                
                # Detect section headers with improved patterns
                section_patterns = [
                    r'^(Section\s+\d+(?:\.\d+)*)\s*[\.\:]\s*(.+)',
                    r'^(Chapter\s+\d+)\s*[\.\:]\s*(.+)',
                    r'^(Article\s+\d+)\s*[\.\:]\s*(.+)',
                    r'^(Part\s+\d+)\s*[\.\:]\s*(.+)',
                    r'^(\d+\.\s*.+)$',
                    r'^([A-Z][A-Z\s]{2,})\s*$'  # All caps headers
                ]
                
                for pattern in section_patterns:
                    section_match = re.match(pattern, line, re.IGNORECASE)
                    if section_match:
                        current_section = section_match.group(0).strip()
                        break
                
                line_size = len(line)
                
                # Check if adding this line would exceed chunk size
                if current_size + line_size > chunk_size and current_chunk:
                    # Create chunk with overlap from previous chunk
                    chunk_content = '\n'.join(current_chunk)
                    
                    chunks.append({
                        "policy_id": policy_id,
                        "content": chunk_content,
                        "content_hash": hashlib.md5(chunk_content.encode()).hexdigest(),
                        "page_number": current_page,
                        "section_title": current_section,
                        "chunk_index": chunk_index,
                        "word_count": len(chunk_content.split()),
                        "char_count": len(chunk_content),
                        "metadata": {
                            "start_page": current_page,
                            "section": current_section,
                            "chunk_type": self._classify_chunk_content(chunk_content)
                        }
                    })
                    
                    # Create overlap for next chunk
                    if overlap > 0 and len(current_chunk) > 1:
                        overlap_lines = current_chunk[-min(3, len(current_chunk)):]  # Last 3 lines for overlap
                        current_chunk = overlap_lines
                        current_size = sum(len(l) for l in overlap_lines)
                    else:
                        current_chunk = []
                        current_size = 0
                    
                    chunk_index += 1
                
                current_chunk.append(line)
                current_size += line_size
            
            # Add the last chunk
            if current_chunk:
                chunk_content = '\n'.join(current_chunk)
                chunks.append({
                    "policy_id": policy_id,
                    "content": chunk_content,
                    "content_hash": hashlib.md5(chunk_content.encode()).hexdigest(),
                    "page_number": current_page,
                    "section_title": current_section,
                    "chunk_index": chunk_index,
                    "word_count": len(chunk_content.split()),
                    "char_count": len(chunk_content),
                    "metadata": {
                        "start_page": current_page,
                        "section": current_section,
                        "chunk_type": self._classify_chunk_content(chunk_content),
                        "is_final_chunk": True
                    }
                })
            
            logger.info(f"Created {len(chunks)} chunks for policy {policy_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking document: {e}")
            raise PolicyParsingException(f"Failed to chunk document: {str(e)}")
    
    def _classify_chunk_content(self, content: str) -> str:
        """Classify chunk content type for better processing"""
        content_lower = content.lower()
        
        # Rule/requirement patterns
        if any(pattern in content_lower for pattern in [
            'shall', 'must', 'required', 'prohibited', 'compliance',
            'limit', 'threshold', 'maximum', 'minimum'
        ]):
            return 'rule_content'
        
        # Definition patterns
        elif any(pattern in content_lower for pattern in [
            'definition', 'means', 'refers to', 'defined as'
        ]):
            return 'definition'
        
        # Procedure patterns
        elif any(pattern in content_lower for pattern in [
            'procedure', 'process', 'step', 'workflow'
        ]):
            return 'procedure'
        
        # Table/data patterns
        elif '|' in content or content.count(':') > 5:
            return 'structured_data'
        
        else:
            return 'general_content'
    
    def get_policy_summary(self, policy_id: str) -> Dict[str, Any]:
        """Get comprehensive summary information about a policy document"""
        try:
            policy = self.db.query(PolicyDocument).filter(
                PolicyDocument.policy_id == policy_id
            ).first()
            
            if not policy:
                return {}
            
            # Get chunk statistics
            chunks = self.db.query(PolicyChunk).filter(
                PolicyChunk.policy_id == policy_id
            ).all()
            
            chunk_stats = {
                "total_chunks": len(chunks),
                "total_words": sum(chunk.word_count or 0 for chunk in chunks),
                "total_chars": sum(chunk.char_count or 0 for chunk in chunks),
                "avg_chunk_size": sum(chunk.word_count or 0 for chunk in chunks) / max(len(chunks), 1),
                "chunk_types": {}
            }
            
            # Analyze chunk types
            for chunk in chunks:
                chunk_type = chunk.metadata.get('chunk_type', 'unknown') if chunk.metadata else 'unknown'
                chunk_stats["chunk_types"][chunk_type] = chunk_stats["chunk_types"].get(chunk_type, 0) + 1
            
            # Get rules count
            from app.models.database import ComplianceRule
            rules_count = self.db.query(ComplianceRule).filter(
                ComplianceRule.source_policy_id == policy_id
            ).count()
            
            return {
                "policy_id": policy.policy_id,
                "filename": policy.filename,
                "document_type": policy.document_type,
                "jurisdiction": policy.jurisdiction,
                "status": policy.status.value if policy.status else None,
                "upload_date": policy.upload_date.isoformat() if policy.upload_date else None,
                "effective_date": policy.effective_date.isoformat() if policy.effective_date else None,
                "expiry_date": policy.expiry_date.isoformat() if policy.expiry_date else None,
                "uploaded_by": policy.uploaded_by,
                "version": policy.version,
                "content_stats": {
                    "content_length": len(policy.content) if policy.content else 0,
                    "content_hash": policy.content_hash
                },
                "chunk_stats": chunk_stats,
                "rules_generated": rules_count,
                "metadata": policy.metadata,
                "created_at": policy.created_at.isoformat() if policy.created_at else None,
                "updated_at": policy.updated_at.isoformat() if policy.updated_at else None
            }
            
        except Exception as e:
            logger.error(f"Error getting policy summary {policy_id}: {e}")
            return {"error": str(e)}
    
    def search_policy_content(self, policy_id: str, search_terms: List[str], 
                            case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """Search for specific terms within a policy document"""
        try:
            chunks = self.db.query(PolicyChunk).filter(
                PolicyChunk.policy_id == policy_id
            ).all()
            
            if not chunks:
                return []
            
            results = []
            
            for chunk in chunks:
                content = chunk.content
                if not case_sensitive:
                    content = content.lower()
                    search_terms = [term.lower() for term in search_terms]
                
                matches = []
                for term in search_terms:
                    if term in content:
                        # Find all occurrences with context
                        start = 0
                        while True:
                            index = content.find(term, start)
                            if index == -1:
                                break
                            
                            # Get context around the match
                            context_start = max(0, index - 100)
                            context_end = min(len(content), index + len(term) + 100)
                            context = content[context_start:context_end]
                            
                            matches.append({
                                "term": term,
                                "position": index,
                                "context": context,
                                "exact_match": content[index:index + len(term)]
                            })
                            
                            start = index + 1
                
                if matches:
                    results.append({
                        "chunk_id": chunk.chunk_id,
                        "chunk_index": chunk.chunk_index,
                        "section_title": chunk.section_title,
                        "page_number": chunk.page_number,
                        "matches": matches,
                        "match_count": len(matches)
                    })
            
            # Sort by relevance (number of matches)
            results.sort(key=lambda x: x["match_count"], reverse=True)
            
            return results
            
        except Exception as e:
            logger.error(f"Error searching policy content: {e}")
            return []
    
    def update_policy_metadata(self, policy_id: str, metadata_updates: Dict[str, Any]) -> bool:
        """Update policy document metadata"""
        try:
            policy = self.db.query(PolicyDocument).filter(
                PolicyDocument.policy_id == policy_id
            ).first()
            
            if not policy:
                return False
            
            # Update allowed fields
            updatable_fields = [
                'document_type', 'jurisdiction', 'effective_date', 'expiry_date',
                'status', 'version'
            ]
            
            for field, value in metadata_updates.items():
                if field in updatable_fields and hasattr(policy, field):
                    if field == 'status' and isinstance(value, str):
                        # Convert string to enum
                        try:
                            value = DocumentStatus(value)
                        except ValueError:
                            continue
                    elif field in ['effective_date', 'expiry_date'] and isinstance(value, str):
                        # Parse date strings
                        try:
                            value = datetime.fromisoformat(value)
                        except ValueError:
                            continue
                    
                    setattr(policy, field, value)
            
            # Update metadata dictionary
            if 'metadata' in metadata_updates:
                if policy.metadata:
                    policy.metadata.update(metadata_updates['metadata'])
                else:
                    policy.metadata = metadata_updates['metadata']
            
            policy.updated_at = datetime.now()
            self.db.commit()
            
            logger.info(f"Updated metadata for policy {policy_id}")
            return True
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Error updating policy metadata: {e}")
            return False
    
    def delete_policy(self, policy_id: str) -> bool:
        """Delete a policy document and all associated data"""
        try:
            policy = self.db.query(PolicyDocument).filter(
                PolicyDocument.policy_id == policy_id
            ).first()
            
            if not policy:
                return False
            
            # Delete associated chunks (cascade should handle this, but explicit is better)
            chunks = self.db.query(PolicyChunk).filter(
                PolicyChunk.policy_id == policy_id
            ).all()
            
            for chunk in chunks:
                self.db.delete(chunk)
            
            # Delete the policy document
            self.db.delete(policy)
            self.db.commit()
            
            logger.info(f"Deleted policy document {policy_id} and {len(chunks)} associated chunks")
            return True
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Error deleting policy {policy_id}: {e}")
            return False
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        try:
            total_policies = self.db.query(PolicyDocument).count()
            total_chunks = self.db.query(PolicyChunk).count()
            
            # Documents by type
            doc_types = self.db.query(
                PolicyDocument.document_type,
                self.db.func.count(PolicyDocument.policy_id)
            ).group_by(PolicyDocument.document_type).all()
            
            # Documents by status
            doc_statuses = self.db.query(
                PolicyDocument.status,
                self.db.func.count(PolicyDocument.policy_id)
            ).group_by(PolicyDocument.status).all()
            
            # Recent uploads (last 30 days)
            thirty_days_ago = datetime.now() - timedelta(days=30)
            recent_uploads = self.db.query(PolicyDocument).filter(
                PolicyDocument.upload_date >= thirty_days_ago
            ).count()
            
            return {
                "total_policies": total_policies,
                "total_chunks": total_chunks,
                "average_chunks_per_policy": round(total_chunks / max(total_policies, 1), 2),
                "document_types": {dt[0]: dt[1] for dt in doc_types if dt[0]},
                "document_statuses": {ds[0].value: ds[1] for ds in doc_statuses if ds[0]},
                "recent_uploads_30_days": recent_uploads,
                "processing_health": "healthy" if total_policies > 0 else "no_data"
            }
            
        except Exception as e:
            logger.error(f"Error getting processing statistics: {e}")
            return {"error": str(e)}